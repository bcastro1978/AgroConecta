import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
import os

def create_element(name):
    return OxmlElement(name)

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=140, bottom=140, left=200, right=200):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table, color="D3D3D3"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(f'''
        <w:tblBorders {nsdecls("w")}>
            <w:top w:val="single" w:sz="4" w:space="0" w:color="{color}"/>
            <w:bottom w:val="single" w:sz="6" w:space="0" w:color="1E3A8A"/>
            <w:left w:val="none"/>
            <w:right w:val="none"/>
            <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{color}"/>
            <w:insideV w:val="none"/>
        </w:tblBorders>
    ''')
    tblPr.append(borders)

def make_callout(doc, text_p_list, bg_hex="F0F4F8", border_hex="1E3A8A"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    
    cell = tbl.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, bg_hex)
    set_cell_margins(cell, top=160, bottom=160, left=220, right=220)
    
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(f'''
        <w:tcBorders {nsdecls("w")}>
            <w:left w:val="single" w:sz="24" w:space="0" w:color="{border_hex}"/>
            <w:top w:val="none"/>
            <w:right w:val="none"/>
            <w:bottom w:val="none"/>
        </w:tcBorders>
    ''')
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    
    # First item in p
    title, desc = text_p_list[0]
    r_title = p.add_run(title + " ")
    r_title.bold = True
    r_title.font.name = "Calibri"
    r_title.font.size = Pt(11)
    r_title.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
    
    r_desc = p.add_run(desc)
    r_desc.font.name = "Calibri"
    r_desc.font.size = Pt(10.5)
    r_desc.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    
    for item in text_p_list[1:]:
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_before = Pt(2)
        p2.paragraph_format.space_after = Pt(4)
        p2.paragraph_format.line_spacing = 1.15
        t, d = item
        rt = p2.add_run(t + " ")
        rt.bold = True
        rt.font.name = "Calibri"
        rt.font.size = Pt(11)
        rt.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
        
        rd = p2.add_run(d)
        rd.font.name = "Calibri"
        rd.font.size = Pt(10.5)
        rd.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def build_architecture_doc():
    doc = Document()

    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Styles setup
    styles = doc.styles
    normal_style = styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    normal_style.paragraph_format.line_spacing = 1.15
    normal_style.paragraph_format.space_after = Pt(6)

    # Document Header / Banner (Title Card)
    banner_table = doc.add_table(rows=1, cols=1)
    banner_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    banner_cell = banner_table.cell(0, 0)
    banner_cell.width = Inches(6.5)
    set_cell_background(banner_cell, "1E3A8A") # Navy Blue
    set_cell_margins(banner_cell, top=240, bottom=240, left=240, right=240)

    p_title = banner_cell.paragraphs[0]
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(4)
    r_t = p_title.add_run("AGROCONECTA")
    r_t.font.name = "Arial"
    r_t.font.size = Pt(26)
    r_t.font.bold = True
    r_t.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    p_sub = banner_cell.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(0)
    r_s = p_sub.add_run("Documento de Arquitectura Tecnológica y Especificaciones del Sistema")
    r_s.font.name = "Arial"
    r_s.font.size = Pt(14)
    r_s.font.color.rgb = RGBColor(0x93, 0xC5, 0xFD) # Soft Blue

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Metadata Block Table
    meta_table = doc.add_table(rows=2, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(meta_table, "E5E7EB")
    
    meta_data = [
        [("Proyecto:", "AgroConecta - Plataforma AgriTech / EUDR"), ("Fecha:", "Agosto 2026")],
        [("Fuente de la Verdad:", "Bóveda de Obsidian (docs/)"), ("Estado:", "Versión 1.0 (Oficial)")]
    ]
    
    for r_idx, row in enumerate(meta_data):
        for c_idx, (label, val) in enumerate(row):
            cell = meta_table.cell(r_idx, c_idx)
            cell.width = Inches(3.25)
            set_cell_background(cell, "F9FAFB" if r_idx % 2 == 0 else "FFFFFF")
            set_cell_margins(cell, top=100, bottom=100, left=140, right=140)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            rl = p.add_run(label + " ")
            rl.bold = True
            rl.font.size = Pt(9.5)
            rl.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
            rv = p.add_run(val)
            rv.font.size = Pt(9.5)
            rv.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)

    doc.add_paragraph().paragraph_format.space_after = Pt(16)

    # Helper function for Headings
    def add_custom_h1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.font.name = "Arial"
        r.font.size = Pt(18)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A) # Navy Blue
        
        # Add bottom border/accent line
        pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="12" w:space="4" w:color="1E3A8A"/></w:pBdr>')
        p._p.get_or_add_pPr().append(pBdr)
        return p

    def add_custom_h2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.font.name = "Arial"
        r.font.size = Pt(14)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0x2D, 0x6A, 0x4F) # Deep Forest Green
        return p

    def add_custom_h3(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.font.name = "Arial"
        r.font.size = Pt(12)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
        return p

    # --- SECCIÓN 1: RESUMEN EJECUTIVO Y VISIÓN ---
    add_custom_h1("1. Resumen Ejecutivo y Visión del Sistema")
    
    p = doc.add_paragraph()
    p.add_run("AgroConecta es una plataforma tecnológica integral de alta precisión diseñada para revolucionar la gestión agronómica, la trazabilidad ambiental y el comercio B2B agrícola. Su arquitectura une la adquisición directa de imágenes satelitales del programa espacial europeo (Copernicus CDSE), la inteligencia artificial conversacional y visionaria multimodal (LangGraph + Gemini Flash) y el cumplimiento normativo internacional para normativas de cero deforestación (EUDR TRACES NT).")
    
    make_callout(doc, [
        ("Pilares Arquitectónicos:", "Despliegue modular en capas, desacoplamiento de frontend y backend, aislamiento estricto de multitenancy vía RLS en PostgreSQL/Supabase, e integración omnicanal WhatsApp/Web para el agricultor.")
    ])

    # --- SECCIÓN 2: ARQUITECTURA DE ALTO NIVEL ---
    add_custom_h1("2. Arquitectura de Alto Nivel y Ecosistema Tecnológico")
    
    doc.add_paragraph("El ecosistema de AgroConecta se estructura en 4 capas clave interoperables:")

    # Table of Layers
    layers_table = doc.add_table(rows=5, cols=3)
    layers_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(layers_table)

    headers = ["Capa Tecnológica", "Componentes Principales", "Responsabilidad Principal"]
    hdr_cells = layers_table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].width = [Inches(1.8), Inches(2.2), Inches(2.5)][i]
        set_cell_background(hdr_cells[i], "1E3A8A")
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=140, right=140)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(title)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(10)

    rows_data = [
        ("Presentation Layer (Frontend)", "React + Vite, Tailwind CSS v4, React-Leaflet, jsPDF", "Single Page Application (SPA) para mapas interactivos geoespaciales, tableros analíticos y generación de certificados EUDR TRACES NT."),
        ("AI Agent Ecosystem (Core IA)", "Python, FastAPI, LangGraph, LangChain, Google Gemini Flash", "Orquestación multi-agente para análisis fitosanitario por visión computacional, recomendaciones hídricas/nutricionales y ruteo omnicanal."),
        ("Backend & Data Layer", "Supabase (PostgreSQL 15+), Row Level Security (RLS), PostGIS", "Persistencia de datos relacionales y GeoJSON (EPSG:4326), aislamiento RLS por usuario/tenant y autenticación JWT securizada."),
        ("Integration & Satellite Layer", "Copernicus CDSE API (Sentinel-1, 2, 3, 5P), n8n, WhatsApp API", "Adquisición de telemetría multiespectral/SAR en tiempo real y pasarela omnicanal automatizada para comunicación con productores.")
    ]

    for r_idx, row in enumerate(rows_data):
        row_cells = layers_table.rows[r_idx + 1].cells
        bg_hex = "F9FAFB" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, text in enumerate(row):
            row_cells[c_idx].width = [Inches(1.8), Inches(2.2), Inches(2.5)][c_idx]
            set_cell_background(row_cells[c_idx], bg_hex)
            set_cell_margins(row_cells[c_idx], top=100, bottom=100, left=140, right=140)
            p = row_cells[c_idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(text)
            r.font.size = Pt(9.5)
            if c_idx == 0:
                r.bold = True
                r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # --- SECCIÓN 3: ARQUITECTURA FRONTEND ---
    add_custom_h1("3. Arquitectura Frontend (SPA Geoespacial)")
    
    add_custom_h2("3.1 Stack Tecnológico y Principios de Diseño")
    p = doc.add_paragraph()
    p.add_run("El portal frontend ha sido diseñado bajo los principios de velocidad, interacción fluida y estética moderna Premium (Dark Mode nativo, componentes responsivos pixel-perfect y glassmorfismo).")

    make_callout(doc, [
        ("Framework Base:", "React 19 + Vite para una compilación ultra-rápida y empaquetado optimizado."),
        ("Motor de Estilos:", "Tailwind CSS v4 con variables CSS personalizadas y diseño UI adaptativo."),
        ("Mapas Geoespaciales:", "React-Leaflet integrado con polígonos GeoJSON (EPSG:4326) e superposición de capas satelitales NDVI/NDMI."),
        ("Generación de Documentos:", "jsPDF optimizado para la emisión client-side de certificados técnicos EUDR.")
    ])

    add_custom_h2("3.2 Estructura Modular de Artefactos")
    doc.add_paragraph("El código cliente mantiene una separación clara de responsabilidades:")
    
    bullets_fe = [
        ("src/components/:", " Reutilizables de interfaz (tablas de telemetría, modales de alerta, tarjetas de insumos B2B)."),
        ("src/pages/:", " Vistas de alto nivel (Dashboard Principal, Monitor Satelital, Gestión de Parcelas, Mercado B2B, Login)."),
        ("src/hooks/:", " Hooks personalizados de lógica de estado como useAuth() para la gestión de sesión JWT con Supabase."),
        ("src/lib/:", " Cliente singleton de conexión Supabase y utilidades globales de formato astronómico y geográfico."),
        ("src/types/:", " Definiciones estrictas TypeScript que garantizan consistencia con la base de datos (Parcela, TelemetryRecord, SmartLead).")
    ]
    for b_title, b_desc in bullets_fe:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_after = Pt(3)
        rt = bp.add_run(b_title)
        rt.bold = True
        rt.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
        rd = bp.add_run(b_desc)

    # --- SECCIÓN 4: BASE DE DATOS Y BACKEND ---
    add_custom_h1("4. Arquitectura de Backend y Modelo de Datos")
    
    add_custom_h2("4.1 Motor de Base de Datos y Seguridad (Supabase PostgreSQL)")
    p = doc.add_paragraph()
    p.add_run("El backend de persistencia se apoya en Supabase PostgreSQL. La seguridad y el acceso a datos están gobernados por políticas estrictas de ")
    r_rls = p.add_run("Row Level Security (RLS)")
    r_rls.bold = True
    p.add_run(", garantizando el aislamiento absoluto entre tenants (Productores y Proveedores B2B).")

    add_custom_h2("4.2 Matriz de Políticas RLS (Row Level Security)")
    
    rls_table = doc.add_table(rows=4, cols=3)
    rls_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(rls_table)

    headers_rls = ["Rol de Usuario", "Permisos sobre Datos", "Mecanismo de Control"]
    for i, title in enumerate(headers_rls):
        cell = rls_table.rows[0].cells[i]
        set_cell_background(cell, "1E3A8A")
        set_cell_margins(cell, top=120, bottom=120, left=140, right=140)
        p = cell.paragraphs[0]
        r = p.add_run(title)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(10)

    rls_data = [
        ("Productores (Producers)", "Acceso exclusivo de lectura y escritura únicamente sobre sus parcelas registadas, telemetría asociada y alertas climáticas.", " auth.uid() == parcel.producer_id "),
        ("Proveedores B2B (Providers)", "Acceso a su perfil corporativo, catálogo de ofertas e inserción de respuestas a Smart Leads recibidos.", " auth.uid() == provider.user_id "),
        ("Sistema (Service Role / IA)", "Bypassea RLS mediante service_role_key para procesar cronjobs masivos de ingestión satelital y actualización de memoria de agentes.", " Supabase Service Key Bypass ")
    ]

    for r_idx, row in enumerate(rls_data):
        row_cells = rls_table.rows[r_idx + 1].cells
        bg_hex = "F9FAFB" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, text in enumerate(row):
            cell = row_cells[c_idx]
            set_cell_background(cell, bg_hex)
            set_cell_margins(cell, top=100, bottom=100, left=140, right=140)
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.font.size = Pt(9.5)
            if c_idx == 0:
                r.bold = True
                r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    add_custom_h2("4.3 Entidades Principales del Dominio")
    
    entities = [
        ("users:", " uuid (PK), role ('Producer', 'Provider', 'Admin'), full_name, phone_number, created_at."),
        ("parcels:", " uuid (PK), producer_id (FK), active_crop, geometry (GeoJSON EPSG:4326), area_ha."),
        ("sat_telemetry:", " uuid (PK), parcel_id (FK), mission ('Sentinel-2', 'Sentinel-1'), ndvi_avg, ndmi_avg, vv_avg, cloud_cover, acquisition_date."),
        ("alerts_events:", " uuid (PK), parcel_id (FK), severity ('Baja', 'Media', 'Alta'), anomaly_type, resolved_status."),
        ("b2b_smart_leads:", " uuid (PK), provider_id (FK), parcel_id (FK), category_match, match_score, status ('Nuevo', 'Contactado', 'Cerrado').")
    ]
    for e_name, e_desc in entities:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_after = Pt(3)
        rt = bp.add_run(e_name)
        rt.bold = True
        rt.font.color.rgb = RGBColor(0x2D, 0x6A, 0x4F)
        rd = bp.add_run(e_desc)

    # --- SECCIÓN 5: ARQUITECTURA DE AGENTES IA ---
    add_custom_h1("5. Orquestación de Inteligencia Artificial (LangGraph Multi-Agente)")
    
    p = doc.add_paragraph()
    p.add_run("La inteligencia agronómica de AgroConecta está estructurada mediante un ecosistema multi-agente en ")
    r_lg = p.add_run("LangGraph / FastAPI")
    r_lg.bold = True
    p.add_run(". Un agente supervisor recibe peticiones desde la Web o WhatsApp (vía webhooks de n8n) y realiza un ruteo semántico dinámico hacia especialistas de dominio.")

    make_callout(doc, [
        ("Supervisor Agent (Router):", "Analiza la intención del usuario, extrae metadatos y consulta herramientas (get_parcel_info, get_latest_copernicus_telemetry)."),
        ("Fitopatólogo (pest_specialist):", "Aplica Visión Computacional con Gemini Flash sobre fotografías de hojas/frutos para identificar plagas/enfermedades y prescribir Manejo Integrado de Plagas (MIP)."),
        ("Especialista en Riego (irrigation_specialist):", "Evalúa series temporales de NDMI. Si el índice cae a < 0.40, genera alertas de estrés hídrico severo."),
        ("Analista Económico (economic_analyst):", "Conecta las necesidades del agricultor con la herramienta get_b2b_providers para canalizar oportunidades de compra de insumos.")
    ])

    add_custom_h3("Gestión de Memoria y Estado Conversacional")
    doc.add_paragraph("La memoria conversacional persiste en Supabase dentro de la tabla agent_memory_state. El número de teléfono del usuario o el ID de usuario web actúa como thread_id, permitiendo retomar diálogos contextuales complejos sin pérdida de antecedente agronómico.")

    # --- SECCIÓN 6: INTEGRACIONES SATELITALES COPERNICUS ---
    add_custom_h1("6. Integraciones Satelitales Copernicus CDSE (ESA)")
    
    doc.add_paragraph("A través de la API de Copernicus Data Space Ecosystem (CDSE), el sistema procesa datos biofísicos de múltiples misiones satelitales:")

    copernicus_table = doc.add_table(rows=6, cols=3)
    copernicus_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(copernicus_table)

    headers_cop = ["Misión Satelital", "Indicadores Extraídos", "Aplicación Agronómica"]
    for i, title in enumerate(headers_cop):
        cell = copernicus_table.rows[0].cells[i]
        set_cell_background(cell, "1E3A8A")
        set_cell_margins(cell, top=120, bottom=120, left=140, right=140)
        p = cell.paragraphs[0]
        r = p.add_run(title)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(10)

    cop_data = [
        ("Sentinel-2 (Óptico 10m)", "NDVI, NDRE, CCC, EVI, LAI, FCOVER", "Salud foliar, clorofila, detección temprana de deficiencia de Nitrógeno y densidad de dosel."),
        ("Sentinel-2 & Sentinel-1 SAR", "NDMI, SSM (Surface Soil Moisture), BSI", "Monitoreo de humedad foliar en copa y humedad superficial del suelo (incluso con nubes)."),
        ("Sentinel-1 SAR (Radar Banda C)", "VV / VH Backscatter, RVI", "Monitoreo estructural continuo e invariable ante cobertura nubosa en zonas tropicales."),
        ("Sentinel-3 SLSTR & ERA5", "LST (Land Surface Temp), ETa, ESI", "Evapotranspiración real (mm/día) y detección de heladas o estrés térmico."),
        ("Sentinel-5P & CAMS", "FAPAR, Aerosol Index / CO", "Absorción de radiación para modelos de rendimiento de cosecha y detección de quemas agrícolas.")
    ]

    for r_idx, row in enumerate(cop_data):
        row_cells = copernicus_table.rows[r_idx + 1].cells
        bg_hex = "F9FAFB" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, text in enumerate(row):
            cell = row_cells[c_idx]
            set_cell_background(cell, bg_hex)
            set_cell_margins(cell, top=100, bottom=100, left=140, right=140)
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.font.size = Pt(9.5)
            if c_idx == 0:
                r.bold = True
                r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # --- SECCIÓN 7: TRAZABILIDAD EUDR ---
    add_custom_h1("7. Cumplimiento Normativo y Trazabilidad EUDR")
    
    p = doc.add_paragraph()
    p.add_run("Para cumplir con la regulación de la Unión Europea sobre Productos Libres de Deforestación (EUDR), AgroConecta integra la verificación satelital de no-deforestación posterior a la fecha de corte oficial (31 de diciembre de 2020).")

    bullets_eudr = [
        ("Geolocalización Polygon GeoJSON:", " Captura de coordenadas exactas para parcelas > 4 hectáreas con precisión WGS84 (EPSG:4326)."),
        ("Análisis Histórico Radar/Óptico:", " Evaluación cruzada del Sentinel-2 y Hansen Global Forest Change para certificar la cobertura forestal preexistente."),
        ("Integración TRACES NT:", " Estructuración del archivo de datos para transmisión directa al sistema TRACES NT de la UE."),
        ("Emisión de Certificados PDF:", " Generación de reportes inviolables descargables desde el frontend con firmas digitales y códigos QR de verificación.")
    ]
    for b_title, b_desc in bullets_eudr:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_after = Pt(3)
        rt = bp.add_run(b_title)
        rt.bold = True
        rt.font.color.rgb = RGBColor(0x2D, 0x6A, 0x4F)
        rd = bp.add_run(b_desc)

    # Output file destination
    output_dir = r"c:\PERSONAL\IA\AGROCONECTA\docs\pdf_entregables"
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, "Arquitectura_Tecnologica_AgroConecta.docx")
    
    doc.save(output_path)
    print(f"Document created successfully at: {output_path}")

if __name__ == "__main__":
    build_architecture_doc()
